from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
IMAGE_DIR = ROOT / ".tomota-studio"
OUTPUT = ROOT / "docs" / "Tomota-Studio-使用说明.docx"

# compact_reference_guide preset
PAGE_WIDTH = 8.5
PAGE_HEIGHT = 11.0
MARGIN = 1.0
CONTENT_WIDTH_IN = 6.5
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B76"
GOLD = "B58B2A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D7DEE7"
GREEN = "256D4A"
CAUTION = "7A5A00"
RISK = "9B1C1C"
WHITE = "FFFFFF"

LATIN_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    font: str = LATIN_FONT,
):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            set_cell_width(cell, width)
            cell.width = Inches(width / 1440)
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = BORDER, size: int = 6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_paragraph_keep(paragraph, *, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    if keep_lines:
        keep = OxmlElement("w:keepLines")
        p_pr.append(keep)


def set_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:pageBreakBefore")
    p_pr.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=MUTED)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(abs_ids, default=0) + 1
    def add_scheme(abs_id: int, fmt: str, text: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "269")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), LATIN_FONT)
        fonts.set(qn("w:hAnsi"), LATIN_FONT)
        fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        first_num = numbering.find(qn("w:num"))
        if first_num is None:
            numbering.append(abstract)
        else:
            numbering.insert(list(numbering).index(first_num), abstract)

    add_scheme(next_abs, "bullet", "•")
    add_scheme(next_abs + 1, "decimal", "%1.")
    return next_abs, next_abs + 1


def create_number_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullet(doc: Document, text: str, bullet_num_id: int, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, color=INK)
    return p


def add_number(doc: Document, text: str, decimal_num_id: int):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    run = p.add_run(text)
    set_run_font(run, color=INK)
    return p


def add_kicker(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    set_paragraph_keep(p, keep_next=True)
    run = p.add_run(text.upper())
    set_run_font(run, size=9, color=GOLD, bold=True)
    run.font.all_caps = True
    run.font.character_spacing = Pt(1.5) if hasattr(run.font, "character_spacing") else None
    return p


def add_title(doc: Document, text: str, *, size=30, after=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    set_paragraph_keep(p, keep_next=True)
    run = p.add_run(text)
    set_run_font(run, size=size, color=NAVY, bold=True)
    return p


def add_subtitle(doc: Document, text: str, *, size=14, after=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=DARK_BLUE)
    return p


def add_body(doc: Document, text: str, *, bold=False, color=INK, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1, *, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        set_page_break_before(p)
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_callout(doc: Document, title: str, body: str, *, accent=BLUE):
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(0)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size=6)
    cell = table.cell(0, 0)
    set_row_cant_split(table.rows[0])
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.5, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    body_run = p2.add_run(body)
    set_run_font(body_run, size=10, color=INK)
    return table


def add_data_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        set_cell_shading(hdr.cells[idx], LIGHT_BLUE)
        p = hdr.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for idx, text in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            run = p.add_run(text)
            set_run_font(run, size=9.3, color=INK, bold=(idx == 0))
    set_table_geometry(table, widths)
    return table


def add_figure(doc: Document, filename: str, caption: str, alt_text: str, *, width=6.25):
    image_path = IMAGE_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size=6)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=60, start=60, bottom=60, end=60)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(6)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=MUTED, italic=True)
    set_paragraph_keep(cap, keep_lines=True)
    return table


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    section.top_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    set_table_geometry(table, [4680, 4680], indent_dxa=0)
    set_no_table_borders(table)
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TOMOTA STUDIO · 使用说明")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("本机版 · 2026.08")
    set_run_font(run, size=8.5, color=MUTED)
    header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    add_kicker(doc, "Operator Guide", align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_title(doc, "Tomota Studio 使用说明", size=30, after=8)
    add_subtitle(doc, "AI 共创 · 严格审查 · 番茄运营", size=15, after=3)
    add_subtitle(doc, "从开一本新书，到发布一章合格正文", size=10.5, after=18)
    add_figure(
        doc,
        "studio-preview.png",
        "图 1　Tomota Studio 首页：作品、规划、流程和平台运营的统一入口",
        "Tomota Studio 首页真实截图，展示作品概览和主要导航入口。",
        width=5.85,
    )
    add_body(
        doc,
        "适用对象：使用 Tomota Studio 管理中文网络小说创作、审查与番茄发布的作者",
        color=MUTED,
        after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_body(doc, "文档版本：2026 年 8 月", color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_quick_start(doc: Document, bullet_num_id: int, decimal_num_id: int):
    add_heading(doc, "1. 先看懂：Studio 负责什么", 1, page_break=True)
    add_body(
        doc,
        "Studio 是一本书的操作台。Antigravity 负责提出内容候选；Tomota 负责保存结构、校验产物、执行审查闸门，并决定章节是否具备发布资格。AI 的候选不会自动覆盖正式内容。",
    )
    add_callout(
        doc,
        "最短使用路径",
        "新建作品 → AI 共创全书方向 → 生成并保存分卷/章纲 → 选择本次章节 → 运行严格流程 → 根据审查反馈修订 → 通过后生成番茄发布批次。",
    )
    add_heading(doc, "主导航", 2)
    add_data_table(
        doc,
        ["入口", "主要用途"],
        [
            ["首页", "查看所有作品、当前进度、待处理问题和平台状态。"],
            ["全书·分卷·章节", "按全书、卷、章三级规划；每一级都可与 AI 交互。"],
            ["写作流程", "选择章节并启动严格状态机；查看任务进度、公开摘要和错误。"],
            ["作品工作区", "阅读正文、保存版本、比较稿件，并查看带原文证据的审查问题。"],
            ["番茄运营", "同步账号作品、预览写入差异、确认发布并对账。"],
            ["设置", "检查 Antigravity、Skill、浏览器会话、存储占用和清理预览。"],
        ],
        [2700, 6660],
    )
    add_heading(doc, "首次使用的五步", 2)
    for text in (
        "进入设置，确认 Antigravity 显示“已认证/可运行”；若需要登录，在可见窗口中由本人完成。",
        "新建作品时只填写读者看得见的资料；本地作品编号由系统自动生成。",
        "先完成全书方向，再完成至少一卷卷纲，最后规划需要写作的章节。",
        "在写作流程中勾选本次要处理的章节，启动后持续观察阶段和审查结果。",
        "章节全部通过后再进入番茄运营，先同步、再预览、即时确认后才写入平台。",
    ):
        add_number(doc, text, decimal_num_id)


def add_new_book(doc: Document, decimal_num_id: int):
    add_heading(doc, "2. 新建作品：先和 AI 把方向聊清楚", 1, page_break=True)
    add_figure(
        doc,
        "studio-new-book-ai.png",
        "图 2　新建作品中的 AI 共创侧栏",
        "新建作品界面的真实截图，右侧打开 AI 共创侧栏，可输入创意并查看候选方案。",
        width=6.23,
    )
    for text in (
        "点击“新建作品”，输入书名、作者名、题材和最初钩子；不再手工填写本地编号。",
        "打开“AI 共创”，告诉 AI 目标读者、题材偏好、不要出现的套路，以及你最想保留的核心冲突。",
        "查看候选内容后可以继续追问，例如“把男主写得不公式化”“不要预设完结章数”“强化悬疑而不是阴森感”。",
        "点击“应用到表单”只会填入当前编辑区；确认无误后，再手动创建作品并进入三级规划。",
    ):
        add_number(doc, text, decimal_num_id)
    add_callout(
        doc,
        "编号由系统管理",
        "本地作品编号用于目录和数据库关联，会根据书名自动生成并处理重名。用户不需要理解或维护它，也不会因为改书名而误改现有目录。",
        accent=GREEN,
    )


def add_three_level_outline(doc: Document, bullet_num_id: int):
    add_heading(doc, "3. 三级大纲：全书、分卷、章节分开规划", 1, page_break=True)
    add_figure(
        doc,
        "studio-planning.png",
        "图 3　全书 · 分卷 · 章节三级规划页",
        "三级规划页真实截图，左侧以树形结构展示全书、卷和章节，右侧显示当前层级的编辑字段。",
        width=4.65,
    )
    add_data_table(
        doc,
        ["层级", "本层要回答的问题"],
        [
            ["全书", "作品卖点、核心冲突、主要人物关系、长期谜团、开放式连载方向。"],
            ["一卷", "本卷目标、主要对抗、人物变化、伏笔推进/兑现、卷末落点和下一卷入口。"],
            ["一章", "目标、阻碍、选择、后果、章末钩子、下一章第一拍和目标字数。"],
        ],
        [1700, 7660],
    )
    add_bullet(doc, "规划窗口不是完结章数：例如“向后规划 5 章”只表示本轮继续设计五章。", bullet_num_id)
    add_bullet(doc, "开放式连载不会强迫整本书在某个固定章节收束。", bullet_num_id)
    add_bullet(doc, "每一卷都有独立卷纲，章节必须归属某卷，避免全书只有平铺的章节列表。", bullet_num_id)


def add_outline_ai(doc: Document, decimal_num_id: int):
    add_heading(doc, "4. 大纲 AI 共创：先预览，再应用，再保存", 1, page_break=True)
    add_figure(
        doc,
        "studio-planning-ai.png",
        "图 4　全书层 AI 共创与候选预览",
        "三级大纲页真实截图，右侧 AI 共创面板展示公开进度、候选内容和继续追问入口。",
        width=4.5,
    )
    for text in (
        "在左侧选中“全书”、某一卷或某一章，AI 会自动按当前层级读取上下文。",
        "输入你的意图，例如“给第一卷设计一个中段反转”“把第三章的阻碍改成知识越界风险”。",
        "AI 返回结构化候选；先检查冲突提醒、缺失项和推荐理由，再继续追问。",
        "满意后点击“应用到表单”；候选仍未进入正式资产。",
        "最后点击页面右上角保存，才会写入该书的三级大纲。",
    ):
        add_number(doc, text, decimal_num_id)
    add_body(
        doc,
        "界面显示正在读取什么、准备生成什么、候选依据、冲突和校验结果，帮助判断是否卡住；不会展示模型隐藏的内部思维链。",
        color=MUTED,
        after=0,
    )


def add_workflow_selection(doc: Document, bullet_num_id: int):
    add_heading(doc, "5. 写作流程：按卷选择本次要处理的章节", 1, page_break=True)
    add_figure(
        doc,
        "studio-workflow.png",
        "图 5　按卷分组的章节选择与流程状态",
        "写作流程页真实截图，章节按卷分组，可勾选本次处理范围并查看已有正文、审查和排期状态。",
        width=4.5,
    )
    add_bullet(doc, "先展开目标卷，再勾选本次要处理的章节；未勾选章节不会被顺带生成。", bullet_num_id)
    add_bullet(doc, "同一本书一次只运行一个 Antigravity 任务，章节严格串行，避免 Canon 互相覆盖。", bullet_num_id)
    add_bullet(doc, "运行中可查看阶段、耗时、公开进度和错误；取消只终止当前任务，不会假装推进状态。", bullet_num_id)
    add_bullet(doc, "任务中断后可重试；如果已有输出，系统会先校验再决定是否重新生成。", bullet_num_id)
    add_callout(
        doc,
        "为什么有正文仍可能显示“待生成/待审查”？",
        "正文文件是否存在、是否通过严格审查、是否被后续修改、是否已排期是四件不同的事。Studio 现在分别显示这些状态，不再把“文件存在”误当成“可发布”。",
        accent=CAUTION,
    )


def add_pipeline(doc: Document):
    add_heading(doc, "6. 严格状态机与修改反馈通道", 1, page_break=True)
    add_body(
        doc,
        "每一章按固定闸门推进。任何审查阶段发现问题，都会回到对应返工点；同章最多五轮，仍未解决就标记 blocked，并停止后续章节。",
    )
    add_data_table(
        doc,
        ["阶段", "完成条件"],
        [
            ["故事圣经", "世界规则、术语、人物知识边界和伏笔账本可用。"],
            ["章节设计", "场景目标、阻碍、动机、选择、后果和承接明确。"],
            ["设计审查", "逻辑、人物与伏笔设计无开放问题。"],
            ["正文生成", "只生成指定章节的当前稿，不触碰 Canon。"],
            ["逻辑审查", "时间线、计数、名词、动机和后果成立。"],
            ["人物与去 AI", "声音可区分、知识不越界、对白有功能、表达自然。"],
            ["伏笔与承接", "线索有对象、有推进/兑现，转场和章末承接有效。"],
            ["无提示冷审", "只凭正文和读者已知事实，也能回答谁、做什么、为什么、指什么。"],
            ["Canon 更新", "只从最终正文提取，并附原文证据。"],
        ],
        [2150, 7210],
    )
    add_heading(doc, "如何给 AI 修改反馈", 2)
    add_body(
        doc,
        "在任务详情或对应大纲层打开反馈输入框，直接写“哪里不对 + 想要什么 + 必须保留什么”。反馈会作为下一次候选或返工输入，不会绕过审查闸门。",
    )
    add_callout(
        doc,
        "推荐反馈格式",
        "位置：第 2 场雨棚下；问题：男主语气与配角相同，对话过密；要求：删掉解释性对白，用动作体现戒备；保留：信件背面的血字和章末钩子。每条正式审查问题还必须包含原文引用、分类、违反规则和修复要求。",
        accent=GREEN,
    )
    add_body(doc, "证据为空的“全项通过”会被判为无效，不能更新 Canon，也不能进入发布批次。", color=MUTED, after=0)


def add_workspace(doc: Document):
    add_heading(doc, "7. 作品工作区：正文存在，不等于已经通过", 1, page_break=True)
    add_figure(
        doc,
        "studio-workspace.png",
        "图 6　作品工作区中的正文与审查证据",
        "作品工作区真实截图，左侧列出正文文件，中间显示章节正文，右侧显示审查证据。",
        width=5.82,
    )
    add_data_table(
        doc,
        ["界面状态", "准确含义"],
        [
            ["已排期", "平台已有明确排期；本地仍要核对正文哈希。"],
            ["审后有修改", "曾经审过，但正文后来发生变化，需要重新走相关审查。"],
            ["已有正文 · 待严格审查", "正文文件已存在，但没有当前版本的完整严格审查报告。"],
            ["已规划待生成", "章纲存在，尚未生成正文文件。"],
            ["已通过", "当前正文哈希已经通过全部闸门，才可加入发布候选。"],
        ],
        [2800, 6560],
    )
    add_body(
        doc,
        "修改正文时点击“保存版本”。保存不会自动宣告通过；只要正文哈希变化，旧审查资格就会失效。右侧问题卡可用于定位原文、理解规则并准备下一轮反馈。",
    )


def add_fanqie(doc: Document, bullet_num_id: int):
    add_heading(doc, "8. 番茄运营：先同步，后预览，写入前即时确认", 1, page_break=True)
    add_figure(
        doc,
        "studio-fanqie.png",
        "图 7　番茄作品运营中心",
        "番茄运营中心真实截图，展示账号会话、平台作品、资料差异、章节候选和只读同步入口。",
        width=4.8,
    )
    add_bullet(doc, "每个番茄账号使用独立可见 Chrome 会话；可在账号选择器中切换，作品列表以当前账号的真实同步结果为准。", bullet_num_id)
    add_bullet(doc, "扫码、验证码和风控由用户本人在可见浏览器完成；Studio 不提供 Cookie、Token、密码或验证码导出。", bullet_num_id)
    add_bullet(doc, "资料、封面或章节写入前先生成差异预览，核对目标作品、标题、字数、顺序和正文哈希。", bullet_num_id)
    add_bullet(doc, "页面变化、断网或结果不确定时先重新同步，不能盲目重复提交。", bullet_num_id)
    add_data_table(
        doc,
        ["操作", "即时确认格式"],
        [
            ["资料或封面", "WRITE <batch-id>"],
            ["批量发布", "PUBLISH <batch-id>"],
            ["单章提交审核", "SUBMIT <batch-id>:<chapter>:<hash12>"],
        ],
        [2600, 6760],
    )
    add_body(
        doc,
        "默认禁止：实名认证、合同/版权、收益、银行卡、提现、税务、密码、手机绑定和账号安全设置不由 Studio 自动操作；删除作品、删除已发布章节和撤回公开内容也不会自动执行。",
        bold=True,
        color=RISK,
        after=0,
    )


def add_settings(doc: Document):
    add_heading(doc, "9. 设置与 Antigravity 故障判断", 1, page_break=True)
    add_figure(
        doc,
        "studio-settings.png",
        "图 8　设置页：生成运行时、Skill、浏览器会话与存储",
        "设置页真实截图，显示 Antigravity 连接状态、Skill 哈希、浏览器会话和存储清理入口。",
        width=5.6,
    )
    add_data_table(
        doc,
        ["现象", "先做什么"],
        [
            ["显示“需要登录”", "点击检测连接；若仍失败，在官方可见窗口完成登录后重试。"],
            ["提示地区不支持", "这是 Antigravity API 的地区限制；网络地区恢复支持后重新检测。"],
            ["permission check failed", "确认任务工作目录是当前作品目录，Prompt 与输出路径都在允许范围内。"],
            ["运行很久没有新日志", "查看公开进度和已用时间；先不要重复点击，必要时取消后幂等重试。"],
            ["旧书状态不对", "刷新项目索引并比较正文文件哈希；修改过的正文必须重新标记审查资格。"],
            ["番茄作品不是自己的", "切换到正确账号，重新执行只读同步；不要把本地作品名当成平台同步结果。"],
        ],
        [2500, 6860],
    )
    add_callout(
        doc,
        "清理规则",
        "工作稿最多保留当前稿和上一稿；其余临时稿进入七天回收区。单书超过 100 MB 时只清理最旧回收内容，最终稿、章纲、Canon、人物/伏笔账本、未解决问题和发布记录不会自动删除。手动清理默认只预览。",
    )


def add_daily_checklist(doc: Document, decimal_num_id: int, bullet_num_id: int):
    add_heading(doc, "10. 一次完整创作的推荐操作顺序", 1, page_break=True)
    for text in (
        "在“全书”层确认这一阶段的故事方向，不把规划窗口误认为完结章数。",
        "在目标卷补齐卷纲：目标、冲突、人物变化、伏笔推进和卷末入口。",
        "在章节层逐章生成/修改章纲，并用 AI 对话追问不清楚的动机、转折和人物声音。",
        "进入写作流程，按卷勾选本轮章节；不要一次选择超出可审查范围的大批章节。",
        "观察 Antigravity 的公开进度；若失败，先看错误属于登录、路径、格式还是内容闸门。",
        "在作品工作区处理带原文证据的问题，保存新版本后重新执行对应审查。",
        "确认冷审、Canon 更新和正文哈希全部通过，再生成番茄发布预览。",
        "切换到正确番茄账号，只读同步；核对作品 ID、章节重复、标题、字数、顺序和哈希。",
        "输入本次批次要求的即时确认，再执行写入；完成后重新同步并对账平台 ID 与状态。",
    ):
        add_number(doc, text, decimal_num_id)

    add_heading(doc, "发布前 30 秒检查", 2)
    for text in (
        "当前选择的是正确的本地作品、番茄账号和平台作品 ID。",
        "章节状态是“已通过”，不是“已有正文”或“审后有修改”。",
        "发布预览里的标题、字数、章节号、排期和 hash12 与当前正文一致。",
        "没有待解决的逻辑、人物、对白、伏笔或冷审问题。",
        "需要公开写入的批次已在当下完成明确确认。",
    ):
        add_bullet(doc, text, bullet_num_id)

    add_callout(
        doc,
        "一句话原则",
        "AI 负责提出候选，作者负责选择与反馈，Tomota 负责验证和留痕；没有通过当前正文的严格闸门，就不进入发布。",
        accent=GREEN,
    )
    add_body(
        doc,
        "—— 完 ——",
        color=MUTED,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def set_core_properties(doc: Document):
    props = doc.core_properties
    props.title = "Tomota Studio 使用说明"
    props.subject = "AI 共创、严格审查与番茄作品运营操作手册"
    props.author = "Tomota Studio"
    props.keywords = "Tomota Studio, Antigravity, 小说写作, 番茄小说, 使用说明"
    props.comments = "基于 2026 年 8 月 Tomota Studio 本机界面制作。"


def audit_document(doc: Document):
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.5
    assert round(section.page_height.inches, 2) == 11.0
    assert all(
        round(value.inches, 2) == 1.0
        for value in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin)
    )
    normal = doc.styles["Normal"]
    assert normal.font.name == LATIN_FONT
    assert round(normal.font.size.pt, 1) == 11.0
    assert round(normal.paragraph_format.space_after.pt, 1) == 6.0
    for filename in (
        "studio-preview.png",
        "studio-new-book-ai.png",
        "studio-planning.png",
        "studio-planning-ai.png",
        "studio-workflow.png",
        "studio-workspace.png",
        "studio-fanqie.png",
        "studio-settings.png",
    ):
        assert (IMAGE_DIR / filename).exists(), filename


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_abs_id, decimal_abs_id = configure_numbering(doc)
    set_core_properties(doc)

    add_cover(doc)
    add_quick_start(
        doc,
        create_number_instance(doc, bullet_abs_id),
        create_number_instance(doc, decimal_abs_id),
    )
    add_new_book(doc, create_number_instance(doc, decimal_abs_id))
    add_three_level_outline(doc, create_number_instance(doc, bullet_abs_id))
    add_outline_ai(doc, create_number_instance(doc, decimal_abs_id))
    add_workflow_selection(doc, create_number_instance(doc, bullet_abs_id))
    add_pipeline(doc)
    add_workspace(doc)
    add_fanqie(doc, create_number_instance(doc, bullet_abs_id))
    add_settings(doc)
    add_daily_checklist(
        doc,
        create_number_instance(doc, decimal_abs_id),
        create_number_instance(doc, bullet_abs_id),
    )

    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
